import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    content: str
    source_file: str
    page: Optional[int] = None
    metadata: dict = field(default_factory=dict)


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def __init__(self):
        pass

    def load_file(self, file_path: str | Path) -> List[DocumentChunk]:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {self.SUPPORTED_EXTENSIONS}")

        logger.info(f"加载文件: {file_path}")

        if ext == ".pdf":
            return self._load_pdf(file_path)
        elif ext == ".docx":
            return self._load_docx(file_path)
        elif ext == ".txt":
            return self._load_txt(file_path)

        return []

    def load_directory(self, dir_path: str | Path) -> List[DocumentChunk]:
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"目录不存在: {dir_path}")

        chunks = []
        for file_path in sorted(dir_path.iterdir()):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    chunks.extend(self.load_file(file_path))
                except Exception as e:
                    logger.error(f"加载文件失败 {file_path}: {e}")
        return chunks

    def _load_pdf(self, file_path: Path) -> List[DocumentChunk]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf")

        chunks = []
        reader = PdfReader(str(file_path))
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text()
                if text and text.strip():
                    chunks.append(DocumentChunk(
                        content=text.strip(),
                        source_file=file_path.name,
                        page=page_num,
                        metadata={"page": page_num, "file": str(file_path)}
                    ))
            except Exception as e:
                logger.warning(f"提取 PDF 第 {page_num} 页失败: {e}")
        logger.info(f"PDF 加载完成: {file_path.name}, 共 {len(chunks)} 页")
        return chunks

    def _load_docx(self, file_path: Path) -> List[DocumentChunk]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        chunks = []
        doc = Document(str(file_path))

        full_text_parts = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                full_text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text_parts.append(" | ".join(row_text))

        full_text = "\n".join(full_text_parts)
        if full_text:
            chunks.append(DocumentChunk(
                content=full_text,
                source_file=file_path.name,
                page=None,
                metadata={"file": str(file_path)}
            ))
        logger.info(f"Word 加载完成: {file_path.name}")
        return chunks

    def _load_txt(self, file_path: Path) -> List[DocumentChunk]:
        try:
            import chardet
        except ImportError:
            raise ImportError("请安装 chardet: pip install chardet")

        raw_data = file_path.read_bytes()
        detection = chardet.detect(raw_data)
        encoding = detection.get("encoding", "utf-8") or "utf-8"

        try:
            text = raw_data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            text = raw_data.decode("utf-8", errors="ignore")

        chunks = []
        if text and text.strip():
            chunks.append(DocumentChunk(
                content=text.strip(),
                source_file=file_path.name,
                page=None,
                metadata={"file": str(file_path), "encoding": encoding}
            ))
        logger.info(f"TXT 加载完成: {file_path.name}")
        return chunks
